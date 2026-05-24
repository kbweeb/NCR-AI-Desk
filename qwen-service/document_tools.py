"""Extract and export office documents for the AI desk."""

from __future__ import annotations

import io
import re
from pathlib import Path

from fpdf import FPDF
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
    ".txt",
    ".md",
    ".markdown",
    ".rtf",
    ".csv",
    ".log",
}


def detect_format(filename: str) -> str:
    ext = Path(filename or "").suffix.lower()
    if ext == ".doc":
        return "docx"  # export target when legacy .doc uploaded as text-only
    if ext in {".md", ".markdown"}:
        return "md"
    if ext in SUPPORTED_EXTENSIONS:
        return ext.lstrip(".")
    return "txt"


def extract_document(data: bytes, filename: str) -> dict:
    ext = Path(filename or "upload.txt").suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext == ".docx":
        return _extract_docx(data)
    if ext == ".doc":
        raise ValueError(
            "Legacy .doc files are not supported. Save as .docx or .pdf and upload again."
        )
    if ext in {".txt", ".md", ".markdown", ".csv", ".log"}:
        return _extract_plain_text(data, filename)
    if ext == ".rtf":
        return _extract_rtf(data, filename)
    raise ValueError(
        f"Unsupported file type '{ext or '(none)'}'. "
        f"Use: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def _extract_pdf(data: bytes) -> dict:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ValueError(
            f"Could not open this PDF ({exc}). "
            "Try re-saving it from Word or printing to PDF again."
        ) from exc
    pages: list[str] = []
    for page in reader.pages:
        pages.append((page.extract_text() or "").strip())
    full = "\n\n".join(p for p in pages if p).strip()
    if not full and len(reader.pages) > 0:
        raise ValueError(
            "No selectable text in this PDF. It may be a scanned form or image-only. "
            "Save as Word (.docx), or export a text-based PDF, then upload again."
        )
    return {
        "text": full,
        "pageCount": len(reader.pages),
        "charCount": len(full),
        "format": "pdf",
    }


def _extract_docx(data: bytes) -> dict:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError(
            "Word support is not installed. Run: pip install python-docx"
        ) from exc

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    full = "\n\n".join(parts).strip()
    return {
        "text": full,
        "pageCount": max(1, len(parts) // 40 + 1),
        "charCount": len(full),
        "format": "docx",
    }


def _extract_plain_text(data: bytes, filename: str) -> dict:
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            text = data.decode(enc).strip()
            break
        except UnicodeDecodeError:
            text = ""
    else:
        text = data.decode("utf-8", errors="replace").strip()
    return {
        "text": text,
        "pageCount": 1,
        "charCount": len(text),
        "format": detect_format(filename),
    }


def _extract_rtf(data: bytes, filename: str) -> dict:
    raw = data.decode("cp1252", errors="replace")
    text = re.sub(r"\\[a-z]+\d* ?", "", raw)
    text = re.sub(r"[{}]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return {
        "text": text,
        "pageCount": 1,
        "charCount": len(text),
        "format": "rtf",
    }


def export_document(title: str, body: str, filename: str, fmt: str | None = None) -> tuple[bytes, str, str]:
    """Returns (bytes, filename, mime_type)."""
    export_fmt = (fmt or detect_format(filename)).lower().replace(".", "")
    stem = Path(filename or "document").stem or "document"
    safe_stem = re.sub(r'[<>:"/\\|?*]+', "_", stem)[:120]

    if export_fmt == "pdf":
        content = _generate_pdf_bytes(title, body)
        name = f"{safe_stem}.pdf"
        return content, name, "application/pdf"
    if export_fmt == "docx":
        content = _generate_docx_bytes(title, body)
        name = f"{safe_stem}.docx"
        return (
            content,
            name,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if export_fmt in {"md", "markdown"}:
        content = body.encode("utf-8")
        name = f"{safe_stem}.md"
        return content, name, "text/markdown; charset=utf-8"
    if export_fmt == "rtf":
        content = _generate_rtf_bytes(title, body)
        name = f"{safe_stem}.rtf"
        return content, name, "application/rtf"
    # txt, csv, log, default
    content = body.encode("utf-8")
    name = f"{safe_stem}.txt"
    return content, name, "text/plain; charset=utf-8"


def _sanitize_for_pdf(text: str) -> str:
    replacements = {
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u2026": "...",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _generate_pdf_bytes(title: str, body: str) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", style="B", size=14)
    pdf.multi_cell(0, 8, _sanitize_for_pdf(title.strip() or "Document"))
    pdf.ln(4)
    pdf.set_font("Helvetica", size=11)
    safe_body = _sanitize_for_pdf(body.strip() or "(empty)")
    for block in re.split(r"\n{2,}", safe_body):
        block = block.strip()
        if not block:
            continue
        if block.startswith(("# ", "## ", "### ")):
            heading = block.lstrip("#").strip()
            pdf.set_font("Helvetica", style="B", size=11)
            pdf.multi_cell(0, 7, heading)
            pdf.set_font("Helvetica", size=11)
            pdf.ln(2)
            continue
        if block.startswith(("- ", "* ")):
            for line in block.splitlines():
                line = line.strip()
                if line.startswith(("- ", "* ")):
                    line = line[2:].strip()
                pdf.multi_cell(0, 6, f"  - {line}")
            pdf.ln(2)
            continue
        pdf.multi_cell(0, 6, block)
        pdf.ln(2)
    out = pdf.output()
    return out if isinstance(out, bytes) else out.encode("latin-1")


def _generate_docx_bytes(title: str, body: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title.strip() or "Document", level=0)
    for block in re.split(r"\n{2,}", body.strip() or ""):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            doc.add_heading(block[2:].strip(), level=1)
        elif block.startswith("## "):
            doc.add_heading(block[3:].strip(), level=2)
        elif block.startswith("### "):
            doc.add_heading(block[4:].strip(), level=3)
        elif block.startswith(("- ", "* ")):
            for line in block.splitlines():
                line = line.strip()
                if line.startswith(("- ", "* ")):
                    doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(block)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_rtf_bytes(title: str, body: str) -> bytes:
    escaped = body.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")
    escaped = escaped.replace("\n", "\\par ")
    rtf = f"{{\\rtf1\\ansi\\deff0{{\\b\\fs28 {title}}}\\par\\par {escaped}}}"
    return rtf.encode("utf-8", errors="replace")


# Backward compatibility
def extract_text_from_pdf(data: bytes) -> dict:
    return _extract_pdf(data)


def generate_pdf_bytes(title: str, body: str) -> bytes:
    content, _, _ = export_document(title, body, "document.pdf", "pdf")
    return content
